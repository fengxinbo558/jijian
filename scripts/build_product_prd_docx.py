"""Build the formal PRD DOCX from the reviewed Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_product_proposal_docx as base


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "product" / "02-机鉴PRD产品需求文档.md"
OUTPUT = ROOT / "docs" / "product" / "02-机鉴PRD产品需求文档.docx"

FONT = "Arial Unicode MS"
INK = "172B3A"
DARK_BLUE = "13465B"
BLUE = "176B87"
TEAL = "1C9A91"
MUTED = "667985"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "EAF5F3"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, size=10.5, color=INK, bold=None, italic=None) -> None:
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().rFonts
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("机鉴 · IDC AI故障调查与现场协同平台  |  PRD V0.1")
    set_run_font(run, 8.5, MUTED)
    p_pr = header._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "D4E1E5")
    borders.append(bottom)
    p_pr.append(borders)
    base.add_page_field(section.footer.paragraphs[0])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    for attr in ("ascii", "hAnsi", "eastAsia"):
        normal._element.rPr.rFonts.set(qn(f"w:{attr}"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        for attr in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_inline(paragraph, text: str, size=10.5, color=INK) -> None:
    token = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)")
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size, color)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_run_font(run, size, color, bold=True)
        elif value.startswith("["):
            label, url = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", value).groups()
            base.add_hyperlink(paragraph, label, url)
        else:
            clean = value.rstrip("。；，、")
            suffix = value[len(clean) :]
            base.add_hyperlink(paragraph, clean, clean)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size, color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size, color)


def add_masthead(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_footer.add_run("内部资料 · 产品评审稿")
    set_run_font(run, 8.5, MUTED)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(7)
    run = kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT  /  V0.1")
    set_run_font(run, 9.5, TEAL, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("机鉴")
    set_run_font(run, 28, INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(3)
    run = subtitle.add_run("IDC AI故障调查与现场协同平台")
    set_run_font(run, 16, DARK_BLUE, bold=True)

    doc_type = doc.add_paragraph()
    doc_type.paragraph_format.space_after = Pt(18)
    run = doc_type.add_run("PRD产品需求文档")
    set_run_font(run, 20, BLUE, bold=True)

    metadata = [
        ("文档编号", "IDC-AI-OPS-PRD-002"),
        ("版本 / 状态", "V0.1 / 产品评审稿"),
        ("编制日期", "2026年8月27日"),
        ("上位文档", "《机鉴产品立项报告》V0.1"),
        ("适用范围", "产品、研发、测试、系统组、网络组、IDC现场、AI运营和试点评审"),
        ("保密等级", "内部资料"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    base.set_table_geometry(table, [1800, 7560])
    table.style = "Table Grid"
    base.set_repeat_table_header(table.rows[0])
    for row_index, (label, value) in enumerate(metadata):
        left, right = table.rows[row_index].cells
        base.set_cell_shading(left, LIGHT_BLUE)
        for cell, text, bold in ((left, label, True), (right, value, False)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(text)
            set_run_font(run, 9.6, DARK_BLUE if bold else INK, bold=bold)

    summary = doc.add_table(rows=1, cols=1)
    base.set_table_geometry(summary, [CONTENT_WIDTH_DXA])
    base.set_repeat_table_header(summary.rows[0])
    base.set_cell_shading(summary.cell(0, 0), LIGHT_TEAL)
    p = summary.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    add_inline(
        p,
        "本PRD同时记录当前已实现能力、真实环境未验证能力和生产前差距；模拟结果不作为生产准确率证明。",
        9.8,
        DARK_BLUE,
    )
    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("目录")
    set_run_font(run, 17, BLUE, bold=True)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run("导入飞书后，可依据Word Heading 1/2/3自动生成可折叠的上下级目录。")
    set_run_font(run, 9.5, MUTED)

    for index, text in enumerate(headings, 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.1)
        paragraph.paragraph_format.space_after = Pt(3)
        number = paragraph.add_run(f"{index:02d}  ")
        set_run_font(number, 9.4, TEAL, bold=True)
        run = paragraph.add_run(text)
        set_run_font(run, 9.7, DARK_BLUE, bold=True)
    doc.add_page_break()


def set_row_cannot_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def table_widths(rows: list[list[str]], columns: int) -> list[int]:
    first = rows[0][0] if rows and rows[0] else ""
    if columns == 4:
        if first == "字段":
            return [1560, 900, 3680, 3220]
        if first == "功能":
            return [1800, 2200, 3200, 2160]
        if first == "一级入口":
            return [1500, 2220, 2160, 3480]
        if first in {"对象", "能力", "情况", "门禁"}:
            return [1500, 2500, 3000, 2360]
        if first == "需求范围":
            return [1600, 2300, 3000, 2460]
        return [1500, 2300, 3200, 2360]
    if columns == 3:
        return [1800, 3780, 3780]
    if columns == 2:
        return [2160, 7200]
    base_width = CONTENT_WIDTH_DXA // columns
    widths = [base_width] * columns
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    base.set_table_geometry(table, table_widths(rows, columns))
    base.set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        set_row_cannot_split(table.rows[row_index])
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = values[col_index] if col_index < len(values) else ""
            if row_index == 0:
                base.set_cell_shading(cell, DARK_BLUE)
            elif row_index % 2 == 0:
                base.set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline(paragraph, text, 8.6, WHITE if row_index == 0 else INK)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def add_requirement_label(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_TEAL)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), TEAL)
    borders.append(left)
    p_pr.append(borders)
    add_inline(paragraph, text, 9.6, DARK_BLUE)


def add_body_paragraph(doc: Document, text: str) -> None:
    if text.startswith("**PRD-") and text.endswith("**"):
        add_requirement_label(doc, text)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    add_inline(paragraph, text)


def add_list_item(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([level, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    add_inline(paragraph, text)


def build_body(doc: Document, lines: list[str]) -> None:
    page_break_heads = {
        "附录A：角色与权限矩阵",
        "附录H：能力真实性状态表",
    }
    index = 0
    paragraph_buffer: list[str] = []
    active_list: str | None = None
    list_num_id: int | None = None

    def flush() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = "".join(item.strip() for item in paragraph_buffer).strip()
            add_body_paragraph(doc, text)
            paragraph_buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush()
            active_list = None
            list_num_id = None
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            flush()
            active_list = None
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue
        if line.startswith("## "):
            flush()
            active_list = None
            text = line[3:].strip()
            if text in page_break_heads:
                doc.add_page_break()
            doc.add_paragraph(text, style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            flush()
            active_list = None
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("#### "):
            flush()
            active_list = None
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            index += 1
            continue
        if line.startswith("- "):
            flush()
            if active_list != "bullet":
                list_num_id = base.add_numbering_definition(doc, True)
                active_list = "bullet"
            add_list_item(doc, line[2:].strip(), list_num_id)
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            flush()
            if active_list != "decimal":
                list_num_id = base.add_numbering_definition(doc, False)
                active_list = "decimal"
            add_list_item(doc, re.sub(r"^\d+\. ", "", line), list_num_id)
            index += 1
            continue
        if line == "---":
            flush()
            active_list = None
            index += 1
            continue
        active_list = None
        list_num_id = None
        paragraph_buffer.append(raw)
        index += 1
    flush()


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 一、"))
    body_lines = lines[start:]
    headings = [line[3:].strip() for line in body_lines if line.startswith("## ")]

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_masthead(doc)
    add_contents(doc, headings)
    build_body(doc, body_lines)

    props = doc.core_properties
    props.title = "机鉴 · IDC AI故障调查与现场协同平台 PRD产品需求文档"
    props.subject = "产品、研发、测试、系统组、网络组、IDC现场、AI运营和试点评审"
    props.author = "机鉴项目组"
    props.keywords = "IDC, AI, PRD, 故障调查, RAG, Agent, 现场协同, OMS"
    props.comments = "V0.1 产品评审稿"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
