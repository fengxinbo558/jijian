"""Build the three remaining core product DOCX documents from Markdown sources."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_product_prd_docx as style
import build_product_proposal_docx as base


ROOT = Path(__file__).resolve().parent.parent
CHINESE_FONT = "Hiragino Sans GB"
style.FONT = CHINESE_FONT
base.FONT = CHINESE_FONT


@dataclass(frozen=True)
class DocSpec:
    source: Path
    output: Path
    english_type: str
    document_type: str
    number: str
    header_label: str
    status: str
    upper: str
    audience: str
    summary: str
    subject: str
    keywords: str


SPECS = (
    DocSpec(
        source=ROOT / "docs/product/03-机鉴技术方案与系统架构设计.md",
        output=ROOT / "docs/product/03-机鉴技术方案与系统架构设计.docx",
        english_type="TECHNICAL ARCHITECTURE  /  V0.1",
        document_type="技术方案与系统架构设计",
        number="IDC-AI-OPS-TECH-003",
        header_label="技术方案 V0.1",
        status="V0.1 / 技术评审稿",
        upper="《机鉴产品立项报告》V0.1、《机鉴PRD产品需求文档》V0.1",
        audience="架构、研发、测试、系统组、网络组、IDC现场、AI运营、安全与试点评审",
        summary="技术边界以当前代码为准：AI负责证据综合，身份、权限、操作许可和审计由确定性系统与人工负责。",
        subject="机鉴技术方案、系统架构、数据流、安全和生产落地差距",
        keywords="IDC, AI, AIOps, 系统架构, RAG, Agent, OMS, 现场协同",
    ),
    DocSpec(
        source=ROOT / "docs/product/04-机鉴AI模型选型报告.md",
        output=ROOT / "docs/product/04-机鉴AI模型选型报告.docx",
        english_type="AI MODEL SELECTION REPORT  /  V0.1",
        document_type="AI模型选型报告",
        number="IDC-AI-OPS-MODEL-004",
        header_label="模型选型 V0.1",
        status="V0.1 / 选型评审稿",
        upper="《机鉴技术方案与系统架构设计》V0.1、《机鉴AI评测方案与基线测试报告》V0.1",
        audience="产品、架构、AI研发、平台研发、安全、采购、合作厂商与试点评审",
        summary="默认本地或私有模型，授权云模型只作对照；最终型号必须通过机鉴盲测，不能按通用榜单直接定标。",
        subject="机鉴生成模型、Embedding、Reranker和供应商适配选型",
        keywords="IDC, AI, 模型选型, Qwen, DeepSeek, GPT, Claude, Gemini, Embedding, Reranker",
    ),
    DocSpec(
        source=ROOT / "docs/product/05-机鉴AI评测方案与基线测试报告.md",
        output=ROOT / "docs/product/05-机鉴AI评测方案与基线测试报告.docx",
        english_type="AI EVALUATION & BASELINE REPORT  /  V0.1",
        document_type="AI评测方案与基线测试报告",
        number="IDC-AI-OPS-EVAL-005",
        header_label="评测基线 V0.1",
        status="V0.1 / 基线评审稿",
        upper="《机鉴技术方案与系统架构设计》V0.1、《机鉴AI模型选型报告》V0.1",
        audience="产品、研发、测试、AI运营、系统组、网络组、IDC现场、安全和试点评审",
        summary="当前通过的是工程、规则、模拟与沙盒基线；真实模型和真实生产准确率必须在只读影子试点中另行验证。",
        subject="机鉴AI评测体系、数据集、指标、硬门禁和当前基线",
        keywords="IDC, AI, 评测, 基线, 盲测, RAG, Agent, 安全门禁, 生产试点",
    ),
)


def set_header(doc: Document, label: str) -> None:
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"机鉴 · IDC AI故障调查与现场协同平台  |  {label}")
    style.set_run_font(run, 8.5, style.MUTED)


def add_masthead(doc: Document, spec: DocSpec) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_footer.add_run("内部资料 · 评审稿")
    style.set_run_font(run, 8.5, style.MUTED)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(7)
    run = kicker.add_run(spec.english_type)
    style.set_run_font(run, 9.5, style.TEAL, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("机鉴")
    style.set_run_font(run, 28, style.INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(3)
    run = subtitle.add_run("IDC AI故障调查与现场协同平台")
    style.set_run_font(run, 16, style.DARK_BLUE, bold=True)

    doc_type = doc.add_paragraph()
    doc_type.paragraph_format.space_after = Pt(18)
    run = doc_type.add_run(spec.document_type)
    style.set_run_font(run, 20, style.BLUE, bold=True)

    metadata = [
        ("文档编号", spec.number),
        ("版本 / 状态", spec.status),
        ("编制日期", "2026年8月27日"),
        ("关联文档", spec.upper),
        ("适用范围", spec.audience),
        ("保密等级", "内部资料"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    base.set_table_geometry(table, [1800, 7560])
    table.style = "Table Grid"
    base.set_repeat_table_header(table.rows[0])
    for row_index, (label, value) in enumerate(metadata):
        left, right = table.rows[row_index].cells
        base.set_cell_shading(left, style.LIGHT_BLUE)
        for cell, text, bold in ((left, label, True), (right, value, False)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(text)
            style.set_run_font(run, 9.3, style.DARK_BLUE if bold else style.INK, bold=bold)

    summary = doc.add_table(rows=1, cols=1)
    base.set_table_geometry(summary, [style.CONTENT_WIDTH_DXA])
    base.set_repeat_table_header(summary.rows[0])
    base.set_cell_shading(summary.cell(0, 0), style.LIGHT_TEAL)
    paragraph = summary.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    style.add_inline(paragraph, spec.summary, 9.6, style.DARK_BLUE)
    doc.add_page_break()


def build_one(spec: DocSpec) -> None:
    text = spec.source.read_text(encoding="utf-8")
    lines = text.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 一、"))
    body_lines = lines[start:]
    headings = [line[3:].strip() for line in body_lines if line.startswith("## ")]

    doc = Document()
    style.configure_page(doc)
    style.configure_styles(doc)
    set_header(doc, spec.header_label)
    add_masthead(doc, spec)
    style.add_contents(doc, headings)
    style.build_body(doc, body_lines)

    props = doc.core_properties
    props.title = f"机鉴 · {spec.document_type}"
    props.subject = spec.subject
    props.author = "机鉴项目组"
    props.keywords = spec.keywords
    props.comments = spec.status
    spec.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(spec.output)
    print(spec.output)


def main() -> None:
    for spec in SPECS:
        build_one(spec)


if __name__ == "__main__":
    main()
