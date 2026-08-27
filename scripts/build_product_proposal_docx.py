"""Build the formal project proposal DOCX from the reviewed Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "product" / "01-机鉴产品立项报告.md"
OUTPUT = ROOT / "docs" / "product" / "01-机鉴产品立项报告.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT = "Arial Unicode MS"
INK = "172B3A"
BLUE = "176B87"
DARK_BLUE = "13465B"
TEAL = "1C9A91"
MUTED = "667985"
LIGHT = "EAF3F5"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11, color=INK, bold=None, italic=None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    run_props.extend([fonts, color, underline, size])
    run.append(run_props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, size=11, color=INK) -> None:
    token = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)")
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size, color)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_run_font(run, size, color, bold=True)
        elif value.startswith("["):
            label, url = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", value).groups()
            add_hyperlink(paragraph, label, url)
        else:
            clean = value.rstrip("。；，、)")
            suffix = value[len(clean) :]
            add_hyperlink(paragraph, clean, clean)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size, color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size, color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering_definition(doc: Document, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, fmt, level_text, justification, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([level, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    add_inline(paragraph, text)


def add_body_paragraph(doc: Document, text: str, callout=False) -> None:
    if callout:
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [CONTENT_WIDTH_DXA])
        set_repeat_table_header(table.rows[0])
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        add_inline(paragraph, text, 11, DARK_BLUE)
        table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.333
    add_inline(paragraph, text)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    if columns == 4:
        widths = [1200, 1560, 4560, 2040]
    elif columns == 3:
        widths = [1800, 3780, 3780]
    elif columns == 2:
        widths = [2160, 7200]
    else:
        base = CONTENT_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            text = values[col_index] if col_index < len(values) else ""
            if row_index == 0:
                set_cell_shading(cell, DARK_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(text)
            set_run_font(run, 9.5, WHITE if row_index == 0 else INK, bold=row_index == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    fp_footer = section.first_page_footer.paragraphs[0]
    fp_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp_footer.add_run("内部资料 · 立项评审稿")
    set_run_font(run, 9, MUTED)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(58)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("IDC AI OPERATIONS")
    set_run_font(run, 11, TEAL, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("机鉴")
    set_run_font(run, 31, INK, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("IDC AI故障调查与现场协同平台")
    set_run_font(run, 17, DARK_BLUE, bold=True)
    subtitle.paragraph_format.space_after = Pt(12)

    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run("产品立项报告")
    set_run_font(run, 22, BLUE, bold=True)
    doc_type.paragraph_format.space_after = Pt(34)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(28)
    p_pr = rule._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), TEAL)
    border.append(bottom)
    p_pr.append(border)

    metadata = [
        ("文档版本", "V0.1"),
        ("文档状态", "立项评审稿"),
        ("编制日期", "2026年8月27日"),
        ("适用范围", "公司内部立项、产品评审、研发与试点准备"),
        ("保密等级", "内部资料"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_geometry(table, [2160, 7200])
    set_repeat_table_header(table.rows[0])
    for index, (label, value) in enumerate(metadata):
        left, right = table.rows[index].cells
        set_cell_shading(left, LIGHT_GRAY)
        for cell, value_text, bold in ((left, label, True), (right, value, False)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value_text)
            set_run_font(run, 10.5, DARK_BLUE if bold else INK, bold=bold)

    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]) -> None:
    heading = doc.add_paragraph("目录", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run("本报告用于立项决策；详细需求、技术、AI评测和部署内容将在后续核心文档中展开。")
    set_run_font(run, 10.5, MUTED)
    for text in headings:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(text)
        set_run_font(run, 10.5, DARK_BLUE, bold=True)
    doc.add_page_break()


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def build_body(doc: Document, lines: list[str]) -> None:
    bullet_num = add_numbering_definition(doc, True)
    decimal_num = add_numbering_definition(doc, False)
    major_breaks = {"七、当前基础与差距", "九、实施路径", "十二、主要风险与控制措施", "十四、参考依据"}
    index = 0
    first_summary = True
    paragraph_buffer: list[str] = []

    def flush() -> None:
        nonlocal paragraph_buffer, first_summary
        if not paragraph_buffer:
            return
        text = "".join(item.strip() for item in paragraph_buffer).strip()
        add_body_paragraph(doc, text, callout=first_summary)
        first_summary = False
        paragraph_buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush()
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            flush()
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue
        if line.startswith("## "):
            flush()
            text = line[3:].strip()
            if text in major_breaks:
                doc.add_page_break()
            doc.add_paragraph(text, style="Heading 1")
            first_summary = text == "一、立项决策摘要"
            index += 1
            continue
        if line.startswith("### "):
            flush()
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if re.match(r"^- ", line):
            flush()
            add_list_item(doc, line[2:].strip(), bullet_num)
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            flush()
            add_list_item(doc, re.sub(r"^\d+\. ", "", line), decimal_num)
            index += 1
            continue
        if line == "---":
            flush()
            index += 1
            continue
        paragraph_buffer.append(raw)
        index += 1
    flush()


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("机鉴 · IDC AI故障调查与现场协同平台  |  产品立项报告")
    set_run_font(run, 8.5, MUTED)
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "D4E1E5")
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 一、"))
    body_lines = lines[start:]
    headings = [line[3:].strip() for line in body_lines if line.startswith("## ")]

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    build_cover(doc)
    add_contents(doc, headings)
    build_body(doc, body_lines)

    props = doc.core_properties
    props.title = "机鉴 · IDC AI故障调查与现场协同平台 产品立项报告"
    props.subject = "公司内部立项、产品评审、研发与试点准备"
    props.author = "机鉴项目组"
    props.keywords = "IDC, AI, 故障调查, RAG, Agent, 现场协同"
    props.comments = "V0.1 立项评审稿"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
